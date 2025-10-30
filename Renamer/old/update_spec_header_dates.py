#!/usr/bin/env python3
import argparse, sys, re, shutil, time, gc
from pathlib import Path
from datetime import datetime

# Dependencies:
#   pip install python-docx pywin32 psutil
from docx import Document

try:
    import win32com.client as win32
    from win32com.client import constants
    import pythoncom
except Exception:
    win32 = None  # only needed for .doc + pdf export

import psutil

DATE_RX = re.compile(
    r"\b("
    r"January|February|March|April|May|June|July|August|September|October|November|December"
    r")\s+([1-9]|[12][0-9]|3[01]),\s+(19|20)\d{2}\b"
)

# Match things like "50% Construction Documents", "90%   Construction    Documents", any percent 0-100, any case
PHASE_RX = re.compile(r"\b(\d{1,3})%\s*Construction\s+Documents\b", re.IGNORECASE)
DEFAULT_PHASE_TEXT = "100% Construction Documents"

# XML namespaces
NS = {
    "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
    "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
    "v": "urn:schemas-microsoft-com:vml",  # legacy shapes sometimes wrap w:t inside
    "wps": "http://schemas.microsoft.com/office/word/2010/wordprocessingShape",
}

def replace_in_all_text_nodes(container_element, repl_fn):
    """
    Replace text in BOTH standard paragraph text (w:t) and shape/textbox text (a:t),
    anywhere inside the given container XML element.
    """
    changed = False
    # All Word text nodes
    for t in container_element.findall(".//w:t", namespaces=NS):
        old = t.text or ""
        new = repl_fn(old)
        if new != old:
            t.text = new
            changed = True

    # All DrawingML text nodes (text inside shapes/textboxes)
    for t in container_element.findall(".//a:t", namespaces=NS):
        old = t.text or ""
        new = repl_fn(old)
        if new != old:
            t.text = new
            changed = True

    return changed

def replace_in_headerlike_anytext(part, target_date, target_phase):
    """
    Works directly on the underlying XML so we also hit shapes/text boxes.
    """
    def repl_fn(text):
        txt = PHASE_RX.sub(target_phase, text)
        txt = DATE_RX.sub(target_date, txt)
        return txt

    # python-docx part -> underlying lxml element is part._element
    return replace_in_all_text_nodes(part._element, repl_fn)

# ---------------------------------------------------------------------
# Utility functions
# ---------------------------------------------------------------------
def format_target_date(date_str: str) -> str:
    cleaned = date_str.replace(",", " ").strip()
    try:
        dt = datetime.strptime(cleaned, "%B %d %Y")
    except ValueError:
        try:
            dt = datetime.strptime(cleaned, "%b %d %Y")
        except ValueError as e:
            raise ValueError("Use 'November 10, 2025' format.") from e
    return dt.strftime("%B %#d, %Y") if sys.platform == "win32" else dt.strftime("%B %-d, %Y")

def iter_all_paragraphs(container):
    for p in getattr(container, "paragraphs", []):
        yield p
    for tbl in getattr(container, "tables", []):
        for row in tbl.rows:
            for cell in row.cells:
                yield from iter_all_paragraphs(cell)

def replace_in_paragraph(paragraph, repl_fn):
    full_text = "".join(run.text for run in paragraph.runs) or paragraph.text
    new_text = repl_fn(full_text)
    if new_text != full_text:
        for _ in range(len(paragraph.runs)):
            r = paragraph.runs[0]
            r.clear(); r.text = ""
            paragraph._p.remove(r._r)
        paragraph.add_run(new_text)
        return True
    return False

def replace_in_headerlike(part, target_date, target_phase):
    """
    Replace (1) date strings like 'November 10, 2025' and
            (2) phase strings like '50% Construction Documents' -> target_phase
    Returns True if any change occurred.
    """
    changed = False

    def repl_fn(text):
        # First: normalize any "...% Construction Documents" to target_phase
        new_text = PHASE_RX.sub(target_phase, text)
        # Then: normalize any long-form Month Day, Year dates
        new_text = DATE_RX.sub(target_date, new_text)
        return new_text

    for p in iter_all_paragraphs(part):
        if replace_in_paragraph(p, repl_fn):
            changed = True
    return changed

def update_docx_dates(path: Path, target_date: str, target_phase: str) -> bool:
    doc = Document(str(path))
    changed = False
    for section in doc.sections:
        # headers
        for hdr in (section.header, section.first_page_header, section.even_page_header):
            if hdr and replace_in_headerlike_anytext(hdr, target_date, target_phase):
                changed = True
        # footers (some templates put the date/phase here)
        for ftr in (section.footer, section.first_page_footer, section.even_page_footer):
            if ftr and replace_in_headerlike_anytext(ftr, target_date, target_phase):
                changed = True
    if changed:
        doc.save(str(path))
    return changed

# ---------------------------------------------------------------------
# Word Automation helpers
# ---------------------------------------------------------------------
def ensure_word():
    if win32 is None:
        raise RuntimeError("pywin32 not installed; cannot handle .doc or PDF export.")
    pythoncom.CoInitialize()
    try:
        app = win32.gencache.EnsureDispatch("Word.Application")
    except Exception:
        app = win32.Dispatch("Word.Application")
    app.Visible = False
    app.DisplayAlerts = 0
    return app

def convert_doc_to_docx(word, doc_path: Path, out_docx: Path):
    out_docx.parent.mkdir(parents=True, exist_ok=True)
    doc = word.Documents.Open(str(doc_path))
    try:
        doc.SaveAs2(str(out_docx), FileFormat=constants.wdFormatXMLDocument)
    finally:
        doc.Close(False)
        gc.collect()

def export_pdf(word, docx_path: Path, pdf_path: Path):
    pdf_path.parent.mkdir(parents=True, exist_ok=True)
    doc = word.Documents.Open(str(docx_path))
    try:
        doc.SaveAs2(str(pdf_path), FileFormat=17)  # 17 = wdFormatPDF
    finally:
        doc.Close(False)
        gc.collect()

def safe_close_word(word):
    try:
        if word:
            word.Quit()
            pythoncom.CoUninitialize()
    except Exception:
        pass
    finally:
        time.sleep(0.3)
        gc.collect()

def kill_orphaned_winword():
    """As a last resort, kill leftover WINWORD processes."""
    for p in psutil.process_iter(["name"]):
        if p.info["name"] and p.info["name"].lower() == "winword.exe":
            p.kill()

# ---------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------
def main():
    p = argparse.ArgumentParser(description="Update header/footer date in Word specs; supports .doc, .docx; optional PDF reprint.")
    p.add_argument("root", help="Folder to scan")
    p.add_argument("--date", "-d", required=True, help="e.g. 'November 10, 2025'")
    p.add_argument("--recursive", "-r", action="store_true", help="Include files in subfolders")
    p.add_argument("--dry-run", action="store_true", help="Show what would change, make no edits")
    p.add_argument("--backup-dir", "-b", default="", help="Copy originals here before editing")
    p.add_argument("--include-doc", action="store_true", help="Process legacy .doc (requires Word)")
    p.add_argument("--replace-doc-inplace", action="store_true",
                   help="After converting .doc->.docx and updating, delete the original .doc and keep the .docx")
    p.add_argument("--reprint-pdf", action="store_true",
                   help="Delete and regenerate same-name PDFs for each updated file (requires Word)")
    p.add_argument("--reprint-pdf-all", action="store_true",
                   help="Regenerate PDFs for all matching files (even if no changes were made). Requires Word.")
    p.add_argument("--set-phase", default=DEFAULT_PHASE_TEXT,
                   help=f"Header phase text to force (default: '{DEFAULT_PHASE_TEXT}')")
    p.add_argument("--exclude-folders", nargs="*", default=["_archive", "archive"],
                   help="Folder names to skip (default: _archive, archive)")
    args = p.parse_args()

    target_date = format_target_date(args.date)
    root = Path(args.root).resolve()
    if not root.exists() or not root.is_dir():
        print(f"Folder not found: {root}"); sys.exit(2)

    backup = Path(args.backup_dir).resolve() if args.backup_dir else None
    if backup:
        backup.mkdir(parents=True, exist_ok=True)

    patterns = ["*.docx"]
    if args.include_doc:
        patterns.append("*.doc")

    files = []
    for pat in patterns:
        iterator = root.rglob(pat) if args.recursive else root.glob(pat)
        for f in iterator:
            # Skip temp and excluded folders
            if f.name.startswith("~$"):
                continue
            if any(part.lower() in [x.lower() for x in args.exclude_folders] for part in f.parts):
                continue
            files.append(f)
    files = sorted(files)

    if not files:
        print("No matching files found."); sys.exit(0)

    print(f"Target date: {target_date}")
    print(f"Scanning {len(files)} file(s)…")

    need_word = args.include_doc or args.reprint_pdf or args.reprint_pdf_all
    word = None
    if need_word:
        try:
            word = ensure_word()
        except Exception as e:
            print(f"[ERROR] Cannot start Word: {e}")
            sys.exit(2)

    updated_ct = 0
    errors = 0

    for f in files:
        try:
            ext = f.suffix.lower()
            if backup and not args.dry_run:
                rel = f.relative_to(root)
                dest = backup / rel
                dest.parent.mkdir(parents=True, exist_ok=True)
                if not dest.exists():
                    shutil.copy2(f, dest)

            if args.dry_run:
                if ext == ".docx":
                    try:
                        doc = Document(str(f))
                        found = False
                        for section in doc.sections:
                            for part in (section.header, section.first_page_header, section.even_page_header,
                                         section.footer, section.first_page_footer, section.even_page_footer):
                                if not part:
                                    continue
                                elem = part._element
                                # Look for either DATE_RX or PHASE_RX in w:t or a:t
                                texts = []
                                texts += [t.text or "" for t in elem.findall(".//w:t", namespaces=NS)]
                                texts += [t.text or "" for t in elem.findall(".//a:t", namespaces=NS)]
                                if any(DATE_RX.search(x) or PHASE_RX.search(x) for x in texts):
                                    found = True
                                    break
                            if found:
                                break
                        if found:
                            print(f"[DRY-RUN] Would update (date/phase): {f}")
                    except Exception as e:
                        print(f"[SKIP] {f} ({e})")
                elif ext == ".doc" and args.include_doc:
                    print(f"[DRY-RUN] Would convert+update: {f}")
                continue

            # real run
            work_docx, original_doc = None, None

            if ext == ".docx":
                work_docx = f
            elif ext == ".doc":
                if not args.include_doc:
                    print(f"[SKIP] (legacy .doc; pass --include-doc) {f}")
                    continue
                work_docx = f.with_suffix(".docx")
                original_doc = f
                convert_doc_to_docx(word, f, work_docx)

            changed = update_docx_dates(work_docx, target_date, args.set_phase or DEFAULT_PHASE_TEXT)
            if changed:
                updated_ct += 1
                print(f"[UPDATED] {f}")

                if original_doc and args.replace_doc_inplace:
                    try: original_doc.unlink(missing_ok=True)
                    except Exception: pass
            else:
                print(f"[NO DATE FOUND] {f}")

            # Handle PDF reprinting (either only when changed, or for all files)
            if args.reprint_pdf or args.reprint_pdf_all:
                # if reprint-pdf (only when changed) OR reprint-pdf-all (always)
                if args.reprint_pdf_all or changed:
                    if word is None:
                        word = ensure_word()
                    pdf_path = work_docx.with_suffix(".pdf")
                    try:
                        if pdf_path.exists():
                            pdf_path.unlink()
                    except Exception:
                        pdf_path.rename(pdf_path.with_suffix(".pdf.bak"))
                    export_pdf(word, work_docx, pdf_path)
                    print(f"  -> [PDF REPRINTED] {pdf_path}")

        except Exception as e:
            errors += 1
            print(f"[ERROR] {f} -> {e}")

    safe_close_word(word)
    kill_orphaned_winword()

    print(f"\nDone. Updated: {updated_ct}, Errors: {errors}")

# ---------------------------------------------------------------------
if __name__ == "__main__":
    main()
