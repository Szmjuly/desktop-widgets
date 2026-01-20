from __future__ import annotations

from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Pt


def _add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def build_user_guide(out_path: Path) -> None:
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    doc.add_heading("Spec Header Date Updater", level=0)
    doc.add_paragraph("User Guide & Manual")
    doc.add_paragraph(f"Generated: {datetime.now().strftime('%B %d, %Y')}")

    doc.add_heading("1. Overview", level=1)
    doc.add_paragraph(
        "Spec Header Date Updater is a Windows tool that batch-updates specification documents by standardizing header/footer "
        "content (date and optional phase text) across a selected folder. It can also regenerate same-name PDFs to keep printed "
        "sets aligned with the updated Word files."
    )

    doc.add_heading("2. What the Tool Updates", level=1)
    _add_bullets(
        doc,
        [
            "Date strings in Word headers/footers (standardized format)",
            "Phase text (optional): patterns like “100% Construction Documents”",
            "Optional: normalize fonts and/or font size (use carefully)",
            "Optional: reprint PDFs from Word so PDFs match updated documents",
        ],
    )

    doc.add_heading("3. Requirements", level=1)
    _add_bullets(
        doc,
        [
            "Windows PC",
            "Access to the target specification folder (read/write permissions)",
            "Microsoft Word installed (required for legacy .doc support and PDF regeneration)",
            "Close Word documents before running to avoid file locks",
        ],
    )

    doc.add_heading("4. Getting Started", level=1)
    doc.add_paragraph("1) Launch the executable.")
    doc.add_paragraph("2) Select the target specification (book) folder.")
    doc.add_paragraph("3) Confirm the Date (and Phase Text if applicable).")
    doc.add_paragraph("4) Choose an Output Mode (recommended default is Update Word docs + reprint PDFs).")
    doc.add_paragraph("5) Click Run Update and review the completion summary/log.")

    doc.add_heading("5. Core Settings", level=1)
    doc.add_heading("5.1 Target Folder", level=2)
    doc.add_paragraph(
        "Select the root folder that contains the specification documents you want to update. The tool can process files in that folder "
        "and (optionally) in subfolders."
    )

    doc.add_heading("5.2 Date", level=2)
    doc.add_paragraph(
        "Choose the date that should appear in document headers/footers. The tool searches for long-form month/day/year dates and replaces "
        "them with the chosen value."
    )

    doc.add_heading("5.3 Phase Text", level=2)
    doc.add_paragraph(
        "Phase Text is optional. It is typically used for a phrase like “100% Construction Documents”. If you leave this field blank, the app "
        "will prompt you to either leave phase text unchanged or explicitly remove it."
    )
    doc.add_paragraph(
        "Important: If you enter a phase format that does not match the expected pattern, the app will warn you because future runs may not be able "
        "to automatically detect and update that phase text."
    )

    doc.add_heading("6. Processing Options", level=1)

    doc.add_heading("6.1 Dry-run", level=2)
    doc.add_paragraph(
        "Dry-run performs a preview scan and reports what would be updated without making changes. Use this first if you are unsure what the tool will change."
    )

    doc.add_heading("6.2 Output Mode", level=2)
    doc.add_paragraph("Output Mode controls whether the run updates Word docs, regenerates PDFs, or both:")
    _add_bullets(
        doc,
        [
            "Update Word docs only: updates .docx/.doc headers/footers but does not change PDFs.",
            "Update Word docs + reprint PDFs: updates docs and regenerates same-name PDFs (recommended when the deliverable is PDFs for printing).",
            "Reprint PDFs only (no doc changes): regenerates PDFs from existing .docx without changing the Word files.",
        ],
    )

    doc.add_heading("6.3 Include subfolders", level=2)
    doc.add_paragraph(
        "When enabled, the tool scans all subfolders under the selected root folder. Disable this if you only want the top-level folder processed."
    )

    doc.add_heading("6.4 Document Handling", level=2)
    doc.add_paragraph("Legacy .doc File Handling:")
    _add_bullets(
        doc,
        [
            "The application automatically detects legacy .doc files before processing",
            "If .doc files are found, you'll be prompted with clear options:",
            "  - Convert to .docx (Delete .doc): Process and remove the original .doc file",
            "  - Convert to .docx (Keep .doc): Process and keep both file versions",
            "  - Skip All .doc Files: Process only .docx files",
            "Word detection: If Microsoft Word is not installed, you'll receive guidance on how to proceed",
            "You can manually enable legacy .doc processing using the checkboxes in Processing Options",
            "Skip 'Table of Contents' files: ignores files with “Table of Contents” in the filename.",
        ],
    )

    doc.add_heading("6.5 Automatic Legacy File Detection", level=2)
    doc.add_paragraph(
        "When you start a processing run, the application automatically scans for legacy .doc files in your target folder."
    )
    doc.add_paragraph("If legacy .doc files are detected:")
    _add_bullets(
        doc,
        [
            "You'll see a dialog showing how many .doc files were found",
            "Clear options explain what will happen with each choice",
            "If Microsoft Word is not installed, you'll receive helpful guidance including:",
            "  - Installing Microsoft Word to enable .doc processing",
            "  - Manually converting .doc files to .docx format",
            "  - Skipping .doc files and processing only .docx files",
        ],
    )
    doc.add_paragraph(
        "This automatic detection prevents accidentally skipping files and ensures you know exactly what will be processed."
    )

    doc.add_heading("7. Backup & Exclusions", level=1)
    doc.add_paragraph(
        "If enabled, the tool can copy files to a backup directory before modifying them. You can also exclude folders by name to skip content like archived sets."
    )

    doc.add_heading("8. Font Normalization (Optional)", level=1)
    doc.add_paragraph(
        "Font normalization changes fonts (and optionally font size) throughout documents. This can affect line breaks, spacing, and pagination."
    )
    _add_bullets(
        doc,
        [
            "Use on a small test set first.",
            "Review output documents for layout changes.",
            "If PDF regeneration is enabled, regenerate PDFs after normalization so PDFs match the updated Word files.",
        ],
    )

    doc.add_heading("9. Activity Log & Results", level=1)
    doc.add_paragraph(
        "The Activity Log shows progress, what files were updated, and any errors encountered. Use Copy to paste results into an email or ticket."
    )

    doc.add_heading("10. Troubleshooting", level=1)
    doc.add_heading("10.1 Word Required / PDF Export Issues", level=2)
    _add_bullets(
        doc,
        [
            "The application automatically detects if Word is installed when .doc files are present",
            "If Word is not available, you'll receive clear guidance on your options",
            "Word is required for: legacy .doc file conversion and PDF regeneration",
            "If you select options that require Word (legacy .doc or PDF regeneration), Word must be installed.",
            "Close any open documents in Word to avoid file lock errors.",
        ],
    )

    doc.add_heading("10.2 Files Not Updating", level=2)
    _add_bullets(
        doc,
        [
            "Check for automatic legacy file detection dialogs - .doc files may have been skipped",
            "If you see a 'Legacy Files Detected' dialog, choose the appropriate conversion option",
            "Run Dry-run to confirm the tool can find date/phase patterns in the documents.",
            "Verify you selected the correct root folder.",
            "If the header text in your files uses a different format than expected, the tool may not match it.",
        ],
    )

    doc.add_heading("10.3 Permissions / Access Errors", level=2)
    _add_bullets(
        doc,
        [
            "Make sure you have write access to the folder.",
            "If files are read-only (or on a protected network share), changes may fail.",
        ],
    )

    doc.add_heading("11. Best Practices", level=1)
    _add_bullets(
        doc,
        [
            "Run Dry-run first when working with a new folder structure.",
            "Use backups for important production sets.",
            "If the deliverable is PDFs, use Update Word docs + reprint PDFs.",
            "Keep Word closed during large runs.",
        ],
    )

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))


def main() -> int:
    project_root = Path(__file__).resolve().parent
    out_path = project_root / "Spec Header Date Updater - User Guide.docx"
    build_user_guide(out_path)
    print(f"Wrote: {out_path}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
