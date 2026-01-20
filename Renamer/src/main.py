#!/usr/bin/env python3
import sys, re, shutil, time, gc, webbrowser, json
from pathlib import Path
from datetime import datetime
from typing import Dict

# ---- third-party deps
from PySide6.QtCore import Qt, QThread, Signal, Slot, QDate, QTimer
from PySide6.QtWidgets import (
    QApplication, QWidget, QGridLayout, QLabel, QLineEdit, QPushButton, QFileDialog,
    QCheckBox, QDateEdit, QTextEdit, QHBoxLayout, QVBoxLayout, QProgressBar,
    QListWidget, QListWidgetItem, QMessageBox, QGroupBox, QDialog, QDialogButtonBox,
    QFormLayout, QFrame, QScrollArea, QSizePolicy, QSpinBox
)
from PySide6.QtGui import QFont, QColor, QPalette

from docx import Document
from docx.shared import Pt
import psutil

# Local imports - Conditional licensing
from src.build_config import INCLUDE_LICENSING

if INCLUDE_LICENSING:
    from src.subscription import SubscriptionManager
else:
    # Dummy class when licensing is disabled
    class SubscriptionManager:
        def __init__(self, *args, **kwargs):
            pass
        def is_subscribed(self):
            return True  # Always return True when licensing disabled
        def get_subscription_info(self):
            return {'status': 'not_included', 'message': 'Licensing not included in this build'}
        def validate_license_key(self, key):
            return True
        def check_document_limit(self):
            return True
        def record_document_processed(self, count=1):
            return True

# pywin32 (optional if .doc or pdf export needed)
try:
    import win32com.client as win32
    from win32com.client import constants
    import pythoncom
except Exception:
    win32 = None  # .doc and PDF export require Word + pywin32

# ----------------------- Modern Design System -----------------------
class Theme:
    """Theme system for light and dark modes."""
    
    LIGHT = {
        "PRIMARY": "#6366F1",      # Indigo 500
        "PRIMARY_HOVER": "#4F46E5",  # Indigo 600
        "SUCCESS": "#10B981",      # Emerald 500
        "DANGER": "#EF4444",       # Red 500
        "WARNING": "#F59E0B",      # Amber 500
        "INFO": "#3B82F6",         # Blue 500
        "BACKGROUND": "#F8FAFC",   # Slate 50
        "CARD": "#FFFFFF",         # Pure white
        "INPUT_BG": "#FFFFFF",     # White inputs
        "BORDER": "#EDF2F7",       # Lighter Slate
        "BORDER_FOCUS": "#6366F1", # Indigo 500
        "TEXT": "#0F172A",         # Slate 900
        "TEXT_SECONDARY": "#64748B",  # Slate 500
        "TEXT_MUTED": "#94A3B8",   # Slate 400
    }
    
    DARK = {
        "PRIMARY": "#818CF8",      # Indigo 400 (lighter for dark mode)
        "PRIMARY_HOVER": "#6366F1",  # Indigo 500
        "SUCCESS": "#34D399",      # Emerald 400
        "DANGER": "#F87171",       # Red 400
        "WARNING": "#FBBF24",      # Amber 400
        "INFO": "#60A5FA",         # Blue 400
        "BACKGROUND": "#0F172A",   # Slate 900
        "CARD": "#1E293B",         # Slate 800
        "INPUT_BG": "#1E293B",     # Slate 800
        "BORDER": "#334155",       # Slate 700
        "BORDER_FOCUS": "#818CF8", # Indigo 400
        "TEXT": "#F1F5F9",         # Slate 100
        "TEXT_SECONDARY": "#CBD5E1",  # Slate 300
        "TEXT_MUTED": "#94A3B8",   # Slate 400
    }
    
    @staticmethod
    def get_theme(is_dark: bool):
        """Get theme colors based on mode."""
        return Theme.DARK if is_dark else Theme.LIGHT


class Colors:
    """Dynamic color system that responds to theme changes."""
    _theme = Theme.LIGHT
    
    @classmethod
    def set_theme(cls, is_dark: bool):
        """Set the current theme."""
        cls._theme = Theme.get_theme(is_dark)
        # Update all color attributes
        cls._update_colors()
    
    @classmethod
    def _update_colors(cls):
        """Update class-level color attributes."""
        cls.PRIMARY = cls._theme["PRIMARY"]
        cls.PRIMARY_HOVER = cls._theme["PRIMARY_HOVER"]
        cls.SUCCESS = cls._theme["SUCCESS"]
        cls.DANGER = cls._theme["DANGER"]
        cls.WARNING = cls._theme["WARNING"]
        cls.INFO = cls._theme["INFO"]
        cls.BACKGROUND = cls._theme["BACKGROUND"]
        cls.CARD = cls._theme["CARD"]
        cls.INPUT_BG = cls._theme["INPUT_BG"]
        cls.BORDER = cls._theme["BORDER"]
        cls.BORDER_FOCUS = cls._theme["BORDER_FOCUS"]
        cls.TEXT = cls._theme["TEXT"]
        cls.TEXT_SECONDARY = cls._theme["TEXT_SECONDARY"]
        cls.TEXT_MUTED = cls._theme["TEXT_MUTED"]
    
    # Initialize with light theme (set initial values directly)
    PRIMARY = Theme.LIGHT["PRIMARY"]
    PRIMARY_HOVER = Theme.LIGHT["PRIMARY_HOVER"]
    SUCCESS = Theme.LIGHT["SUCCESS"]
    DANGER = Theme.LIGHT["DANGER"]
    WARNING = Theme.LIGHT["WARNING"]
    INFO = Theme.LIGHT["INFO"]
    BACKGROUND = Theme.LIGHT["BACKGROUND"]
    CARD = Theme.LIGHT["CARD"]
    INPUT_BG = Theme.LIGHT["INPUT_BG"]
    BORDER = Theme.LIGHT["BORDER"]
    BORDER_FOCUS = Theme.LIGHT["BORDER_FOCUS"]
    TEXT = Theme.LIGHT["TEXT"]
    TEXT_SECONDARY = Theme.LIGHT["TEXT_SECONDARY"]
    TEXT_MUTED = Theme.LIGHT["TEXT_MUTED"]


class ModernCard(QFrame):
    """Modern card widget with shadow effect."""
    def __init__(self, title=None, parent=None):
        super().__init__(parent)
        self.title_label = None
        self._title = title
        self.update_style()
        
        layout = QVBoxLayout(self)
        layout.setSpacing(4)
        layout.setContentsMargins(0, 0, 0, 0)
        
        if title:
            self.title_label = QLabel(title)
            layout.addWidget(self.title_label)
        
        self.content_layout = QVBoxLayout()
        self.content_layout.setSpacing(4)
        layout.addLayout(self.content_layout)
    
    def update_style(self):
        """Update card styling for current theme."""
        self.setStyleSheet(f"""
            QFrame {{
                background-color: {Colors.CARD};
                border: 1px solid {Colors.BORDER};
                border-radius: 6px;
                padding: 6px;
            }}
        """)
        if self.title_label:
            self.title_label.setStyleSheet(f"""
                font-size: 12px;
                font-weight: 600;
                color: {Colors.TEXT};
                padding-bottom: 4px;
                border-bottom: 1px solid {Colors.BORDER};
            """)
    
    def add_widget(self, widget):
        """Add widget to card content."""
        self.content_layout.addWidget(widget)


class ModernButton(QPushButton):
    """Styled button with different variants."""
    def __init__(self, text, variant="primary", icon=None, parent=None):
        super().__init__(text, parent)
        self._variant = variant
        self._icon = icon
        self._original_text = text
        if icon:
            self.setText(f"{icon} {text}")
        else:
            self.setText(text)
        self.update_style()
        self.setCursor(Qt.PointingHandCursor)
    
    def update_style(self):
        """Update button styling for current theme."""
        colors = {
            "primary": (Colors.PRIMARY, "#FFFFFF"),
            "success": (Colors.SUCCESS, "#FFFFFF"),
            "danger": (Colors.DANGER, "#FFFFFF"),
            "secondary": (Colors.BORDER, Colors.TEXT),
        }
        
        bg_color, text_color = colors.get(self._variant, colors["primary"])
        
        disabled_bg = "#64748B" if Colors._theme == Theme.DARK else "#D1D5DB"
        disabled_text = "#94A3B8" if Colors._theme == Theme.DARK else "#9CA3AF"
        
        self.setStyleSheet(f"""
            QPushButton {{
                background-color: {bg_color};
                color: {text_color};
                border: none;
                border-radius: 4px;
                padding: 0px 8px;
                font-size: 11px;
                font-weight: 600;
                height: 26px;
            }}
            QPushButton:hover {{
                opacity: 0.9;
                background-color: {bg_color};
            }}
            QPushButton:pressed {{
                opacity: 0.8;
            }}
            QPushButton:disabled {{
                background-color: {disabled_bg};
                color: {disabled_text};
            }}
        """)


# ----------------------- Core logic (from your script) -----------------------
DATE_RX = re.compile(
    r"\b("
    r"January|February|March|April|May|June|July|August|September|October|November|December"
    r")\s+([1-9]|[12][0-9]|3[01]),\s+(19|20)\d{2}\b"
)

# Match things like "50% Construction Documents", "90%   Construction    Documents", any percent 0-100, any case
PHASE_RX = re.compile(r"\b(\d{1,3})%\s*Construction\s+Documents\b", re.IGNORECASE)
DEFAULT_PHASE_TEXT = "100% Construction Documents"
DEFAULT_FONT_NAME = "Arial"  # Standard font for normalization
DEFAULT_FONT_SIZE = 10  # Default font size in points

def format_target_date(date_str: str) -> str:
    cleaned = date_str.replace(",", " ").strip()
    try:
        dt = datetime.strptime(cleaned, "%B %d %Y")
    except ValueError:
        try:
            dt = datetime.strptime(cleaned, "%b %d %Y")
        except ValueError as e:
            raise ValueError("Use 'November 10, 2025' format.") from e
    return dt.strftime("%B %#d, %Y") if sys.platform == "win32" else dt.strftime("%B %-d, %Y")

def iter_all_paragraphs(container):
    for p in getattr(container, "paragraphs", []):
        yield p
    for tbl in getattr(container, "tables", []):
        for row in tbl.rows:
            for cell in row.cells:
                yield from iter_all_paragraphs(cell)

def replace_in_paragraph(paragraph, repl_fn):
    full_text = "".join(run.text for run in paragraph.runs) or paragraph.text
    new_text = repl_fn(full_text)
    if new_text != full_text:
        for _ in range(len(paragraph.runs)):
            r = paragraph.runs[0]
            r.clear(); r.text = ""
            paragraph._p.remove(r._r)
        paragraph.add_run(new_text)
        return True
    return False

def replace_in_headerlike(part, target_date, target_phase=None):
    """
    Replace (1) date strings like 'November 10, 2025' and
            (2) phase strings like '50% Construction Documents' -> target_phase
    Returns (date_changed, phase_changed) tuple.
    """
    date_changed = False
    phase_changed = False

    def repl_fn_date(text):
        # Normalize any long-form Month Day, Year dates
        new_text = DATE_RX.sub(target_date, text)
        return new_text != text, new_text
    
    def repl_fn_phase(text):
        new_text = PHASE_RX.sub(target_phase, text)
        return new_text != text, new_text

    for p in iter_all_paragraphs(part):
        txt = "".join(run.text for run in p.runs) if hasattr(p, 'runs') else p.text
        if DATE_RX.search(txt):
            date_upd, new_text_date = repl_fn_date(txt)
            if date_upd:
                for run in p.runs if hasattr(p, 'runs') else []:
                    run.text = DATE_RX.sub(target_date, run.text)
                date_changed = True
        if target_phase is not None and PHASE_RX.search(txt):
            phase_upd, new_text_phase = repl_fn_phase(txt)
            if phase_upd:
                for run in p.runs if hasattr(p, 'runs') else []:
                    run.text = PHASE_RX.sub(target_phase, run.text)
                phase_changed = True
    
    return (date_changed, phase_changed)

def normalize_whitespace_in_header_footer(container) -> bool:
    """
    Trim only TRAILING whitespace from the LAST run in each paragraph.
    This prevents line wrapping when font size increases, while preserving
    intentional spacing between elements (like title and page numbers).
    Returns True if any changes were made.
    """
    changed = False
    
    for p in getattr(container, 'paragraphs', []):
        # Only trim trailing whitespace from the LAST run with content
        # Work backwards through runs
        for i in range(len(p.runs) - 1, -1, -1):
            run = p.runs[i]
            if run.text:
                if run.text.strip() == '':
                    # This run is all whitespace at the end - clear it
                    run.text = ''
                    changed = True
                elif run.text.endswith(' ') or run.text.endswith('\t'):
                    # Has trailing whitespace - trim it
                    run.text = run.text.rstrip()
                    changed = True
                    break  # Stop after trimming the last content run
                else:
                    break  # Last run has no trailing whitespace, we're done
    
    # Handle tables in header/footer
    for tbl in getattr(container, 'tables', []):
        for row in tbl.rows:
            for cell in row.cells:
                if normalize_whitespace_in_header_footer(cell):
                    changed = True
    
    return changed


def normalize_fonts_in_document(doc, target_font: str, target_size: int = None, 
                                 trim_header_footer_whitespace: bool = True) -> bool:
    """
    Normalize all fonts in a document to the target font.
    Optionally also sets all fonts to target_size (in points).
    When changing font size, also trims excess whitespace from headers/footers to prevent wrapping.
    Returns True if any changes were made.
    """
    changed = False
    target_size_pt = Pt(target_size) if target_size else None
    
    def normalize_runs(container, is_header_footer=False):
        """Normalize fonts in runs within a container (paragraph, table, etc.)"""
        nonlocal changed
        for p in getattr(container, 'paragraphs', []):
            for run in p.runs:
                # Skip empty runs
                if not run.text:
                    continue
                # Check font name - None means inherited from style, so we should set it explicitly
                current_font = run.font.name
                if current_font is None or current_font != target_font:
                    run.font.name = target_font
                    changed = True
                # Check font size if target specified
                if target_size_pt:
                    current_size = run.font.size
                    if current_size is None or current_size != target_size_pt:
                        run.font.size = target_size_pt
                        changed = True
        for tbl in getattr(container, 'tables', []):
            for row in tbl.rows:
                for cell in row.cells:
                    normalize_runs(cell, is_header_footer)
    
    # Normalize body content
    normalize_runs(doc)
    
    # Normalize headers and footers
    for section in doc.sections:
        for hdr in (section.header, section.first_page_header, section.even_page_header):
            if hdr:
                normalize_runs(hdr, is_header_footer=True)
                # Trim whitespace in headers when changing font size
                if target_size_pt and trim_header_footer_whitespace:
                    if normalize_whitespace_in_header_footer(hdr):
                        changed = True
        for ftr in (section.footer, section.first_page_footer, section.even_page_footer):
            if ftr:
                normalize_runs(ftr, is_header_footer=True)
                # Trim whitespace in footers when changing font size
                if target_size_pt and trim_header_footer_whitespace:
                    if normalize_whitespace_in_header_footer(ftr):
                        changed = True
    
    return changed


def update_docx_dates(path: Path, target_date: str, target_phase=None, 
                      normalize_fonts: bool = False, target_font: str = DEFAULT_FONT_NAME,
                      target_font_size: int = None) -> Dict[str, bool]:
    """
    Update dates and phase text in document.
    Optionally normalize all fonts to target_font and/or target_font_size.
    Returns dict with 'changed', 'date_changed', 'phase_changed', 'fonts_changed'
    """
    doc = Document(str(path))
    date_changed = False
    phase_changed = False
    fonts_changed = False
    changed = False
    
    for section in doc.sections:
        for hdr in (section.header, section.first_page_header, section.even_page_header):
            if hdr:
                date_upd, phase_upd = replace_in_headerlike(hdr, target_date, target_phase)
                if date_upd: date_changed = True
                if phase_upd: phase_changed = True
                if date_upd or phase_upd: changed = True
        for ftr in (section.footer, section.first_page_footer, section.even_page_footer):
            if ftr:
                date_upd, phase_upd = replace_in_headerlike(ftr, target_date, target_phase)
                if date_upd: date_changed = True
                if phase_upd: phase_changed = True
                if date_upd or phase_upd: changed = True
    
    # Font normalization (if enabled)
    if normalize_fonts:
        fonts_changed = normalize_fonts_in_document(doc, target_font, target_font_size)
        if fonts_changed:
            changed = True
    
    if changed:
        doc.save(str(path))
    
    return {
        'changed': changed,
        'date_changed': date_changed,
        'phase_changed': phase_changed,
        'fonts_changed': fonts_changed
    }

def check_word_available() -> tuple[bool, str]:
    """Check if Microsoft Word is available. Returns (is_available, reason)."""
    if win32 is None:
        return False, "pywin32 not installed"
    
    try:
        # Just check if Word COM class exists without starting it
        import winreg
        try:
            key = winreg.OpenKey(winreg.HKEY_CLASSES_ROOT, r"Word.Application\CLSID")
            winreg.CloseKey(key)
            return True, ""
        except WindowsError:
            return False, "Microsoft Word not installed"
    except Exception as e:
        return False, f"Unable to detect Word: {str(e)}"

def ensure_word():
    if win32 is None:
        raise RuntimeError("pywin32 not installed; cannot handle .doc or PDF export.")
    pythoncom.CoInitialize()
    try:
        app = win32.gencache.EnsureDispatch("Word.Application")
    except Exception:
        app = win32.Dispatch("Word.Application")
    app.Visible = False
    app.DisplayAlerts = 0
    return app

def convert_doc_to_docx(word, doc_path: Path, out_docx: Path):
    out_docx.parent.mkdir(parents=True, exist_ok=True)
    doc = word.Documents.Open(str(doc_path))
    try:
        doc.SaveAs2(str(out_docx), FileFormat=constants.wdFormatXMLDocument)
    finally:
        doc.Close(False)
        gc.collect()

def export_pdf(word, docx_path: Path, pdf_path: Path):
    pdf_path.parent.mkdir(parents=True, exist_ok=True)
    doc = word.Documents.Open(str(docx_path))
    try:
        doc.SaveAs2(str(pdf_path), FileFormat=17)  # 17 = wdFormatPDF
    finally:
        doc.Close(False)
        gc.collect()

def export_pdf_fast(word, docx_path: Path, pdf_path: Path):
    """Fast PDF export without gc.collect() - use for batch operations"""
    pdf_path.parent.mkdir(parents=True, exist_ok=True)
    doc = word.Documents.Open(str(docx_path), ReadOnly=True)
    try:
        doc.SaveAs2(str(pdf_path), FileFormat=17)  # 17 = wdFormatPDF
    finally:
        doc.Close(False)

def safe_close_word(word):
    try:
        if word:
            word.Quit()
            pythoncom.CoUninitialize()
    except Exception:
        pass
    finally:
        time.sleep(0.2)
        gc.collect()

def kill_orphaned_winword():
    for p in psutil.process_iter(["name"]):
        if p.info["name"] and p.info["name"].lower() == "winword.exe":
            try:
                if p.username():  # just to touch info and ensure process handle OK
                    pass
            except Exception:
                pass
            # Do NOT kill users' interactive sessions by default; leave disabled.
            # p.kill()
            # If you want a hard kill as last resort, enable the line above.


# ----------------------- Worker (QThread) -----------------------
class UpdateWorker(QThread):
    log = Signal(str)               # plain text log
    progress = Signal(int, int)     # current, total
    finished = Signal(int, int, dict)  # updated_ct, errors, stats_dict
    needsWord = Signal(str)         # error if Word/pywin32 missing
    enableUI = Signal(bool)

    def __init__(self, root: str, date_str: str, phase_text: str|None, recursive: bool, dry_run: bool,
                 backup_dir: str|None, include_doc: bool, replace_doc_inplace: bool,
                 reprint_pdf: bool, exclude_folders: list[str], subscription_mgr=None,
                 default_backup_dir: str|None = None,
                 normalize_fonts: bool = False, target_font: str = DEFAULT_FONT_NAME,
                 target_font_size: int = None, skip_toc: bool = True,
                 reprint_only: bool = False):
        super().__init__()
        self.root = Path(root)
        self.date_str = date_str
        self.phase_text = phase_text
        self.recursive = recursive
        self.dry_run = dry_run
        self.backup_dir = Path(backup_dir) if backup_dir else None
        self.default_backup_dir = Path(default_backup_dir) if default_backup_dir else None
        self.include_doc = include_doc
        self.replace_doc_inplace = replace_doc_inplace
        self.reprint_pdf = reprint_pdf
        self.reprint_only = reprint_only
        self.exclude_folders = [x.lower() for x in exclude_folders]
        self.subscription_mgr = subscription_mgr
        self.normalize_fonts = normalize_fonts
        self.target_font = target_font
        self.target_font_size = target_font_size
        self.skip_toc = skip_toc
        self._cancel = False
        self.start_time = None
        
        # Processing statistics
        self.stats = {
            'files_scanned': 0,
            'documents_updated': 0,
            'documents_with_date_changes': 0,
            'documents_with_phase_changes': 0,
            'documents_with_font_changes': 0,
            'documents_with_both': 0,
            'pdfs_created': 0,
            'errors': 0,
            'duration_seconds': 0.0
        }

    def cancel(self):
        self._cancel = True

    def run(self):
        import time
        self.start_time = time.time()
        self.enableUI.emit(False)
        try:
            target_date = format_target_date(self.date_str)
        except Exception as e:
            self.log.emit(f"[ERROR] {e}")
            self.enableUI.emit(True)
            return

        if not self.root.exists() or not self.root.is_dir():
            self.log.emit(f"[ERROR] Folder not found: {self.root}")
            self.enableUI.emit(True)
            return

        if self.backup_dir:
            self.backup_dir.mkdir(parents=True, exist_ok=True)

        patterns = ["*.docx"]
        if self.include_doc:
            patterns.append("*.doc")

        # gather files
        files: list[Path] = []
        for pat in patterns:
            iterator = self.root.rglob(pat) if self.recursive else self.root.glob(pat)
            for f in iterator:
                if self._cancel: break
                if f.name.startswith("~$"):   # skip Word lock temp files
                    continue
                # skip excluded folders anywhere in path
                if any(part.lower() in self.exclude_folders for part in f.parts):
                    continue
                # skip Table of Contents files if option enabled
                if self.skip_toc and "table of contents" in f.stem.lower():
                    continue
                files.append(f)
        files.sort()

        if not files:
            self.log.emit("No matching files found.")
            self.enableUI.emit(True)
            return

        self.log.emit(f"Target date: {target_date}")
        if self.reprint_only:
            self.log.emit("Mode: REPRINT PDFs ONLY (no document changes)")
        elif self.normalize_fonts:
            size_info = f", size: {self.target_font_size}pt" if self.target_font_size else ""
            self.log.emit(f"Font normalization: ON (font: {self.target_font}{size_info})")
        self.log.emit(f"Scanning {len(files)} file(s)…")

        need_word = self.include_doc or self.reprint_pdf or self.reprint_only
        word = None
        if need_word:
            try:
                word = ensure_word()
            except Exception as e:
                self.needsWord.emit(str(e))
                self.enableUI.emit(True)
                return

        # Check document limit before processing (if licensing enabled)
        if INCLUDE_LICENSING and self.subscription_mgr:
            limit_check = self.subscription_mgr.check_document_limit(requested_count=len(files))
            if not limit_check.get('allowed', True):
                remaining = limit_check.get('remaining', 0)
                limit = limit_check.get('limit', 0)
                self.log.emit(f"[ERROR] Document limit exceeded. Limit: {limit}, Remaining: {remaining}, Requested: {len(files)}")
                self.enableUI.emit(True)
                return
        
        updated_ct = 0
        errors = 0

        # FAST PATH: Reprint-only mode - batch delete then batch export
        if self.reprint_only:
            self.log.emit("Phase 1: Deleting existing PDFs...")
            deleted_count = 0
            docx_files = []
            
            # Batch delete all PDFs first (very fast)
            for f in files:
                if self._cancel:
                    break
                if f.suffix.lower() == ".docx":
                    pdf_path = f.with_suffix(".pdf")
                    if pdf_path.exists():
                        try:
                            pdf_path.unlink()
                            deleted_count += 1
                        except Exception:
                            try:
                                pdf_path.rename(pdf_path.with_suffix(".pdf.bak"))
                                deleted_count += 1
                            except Exception:
                                pass
                    docx_files.append(f)
            
            self.log.emit(f"Deleted {deleted_count} existing PDFs")
            
            if self._cancel:
                self.log.emit("[CANCELLED] Stopping at user request.")
                self.enableUI.emit(True)
                return
            
            self.log.emit(f"Phase 2: Exporting {len(docx_files)} PDFs...")
            
            # Batch export all PDFs (optimized)
            for idx, f in enumerate(docx_files, start=1):
                if self._cancel:
                    self.log.emit("[CANCELLED] Stopping at user request.")
                    break
                try:
                    pdf_path = f.with_suffix(".pdf")
                    export_pdf_fast(word, f, pdf_path)
                    self.stats['pdfs_created'] += 1
                    self.log.emit(f"[{idx}/{len(docx_files)}] {f.name}")
                    self.progress.emit(idx, len(docx_files))
                    
                    # Periodic garbage collection every 10 files
                    if idx % 10 == 0:
                        gc.collect()
                except Exception as e:
                    errors += 1
                    self.log.emit(f"[ERROR] {f.name}: {e}")
            
            # Final cleanup
            gc.collect()
            
            import time
            duration = time.time() - self.start_time
            self.stats['duration_seconds'] = duration
            self.stats['files_scanned'] = len(files)
            self.stats['errors'] = errors
            
            self.log.emit(f"\nCompleted in {duration:.1f}s - {self.stats['pdfs_created']} PDFs created")
            self.finished.emit(self.stats['pdfs_created'], errors, self.stats)
            self.enableUI.emit(True)
            return

        for idx, f in enumerate(files, start=1):
            if self._cancel:
                self.log.emit("[CANCELLED] Stopping at user request.")
                break
            
            # Periodic limit check during processing (every 10 files)
            if INCLUDE_LICENSING and self.subscription_mgr and idx % 10 == 0:
                limit_check = self.subscription_mgr.check_document_limit(requested_count=1)
                if not limit_check.get('allowed', True):
                    remaining = limit_check.get('remaining', 0)
                    self.log.emit(f"[WARNING] Document limit reached. Processed: {updated_ct}, Remaining: {remaining}")
                    break

            try:
                ext = f.suffix.lower()

                if self.backup_dir and not self.dry_run:
                    rel = f.relative_to(self.root)
                    dest = self.backup_dir / rel
                    dest.parent.mkdir(parents=True, exist_ok=True)
                    if not dest.exists():
                        shutil.copy2(f, dest)

                if self.dry_run:
                    if ext == ".docx":
                        try:
                            doc = Document(str(f))
                            found = False
                            for section in doc.sections:
                                # check header+footer containers
                                for part in (section.header, section.first_page_header, section.even_page_header,
                                             section.footer, section.first_page_footer, section.even_page_footer):
                                    if not part: continue
                                    for pgraph in iter_all_paragraphs(part):
                                        txt = "".join(run.text for run in pgraph.runs) or pgraph.text
                                        if DATE_RX.search(txt) or PHASE_RX.search(txt):
                                            found = True; break
                                if found: break
                            if found: self.log.emit(f"[DRY-RUN] Would update (date/phase): {f}")
                        except Exception as e:
                            self.log.emit(f"[SKIP] {f} ({e})")
                    elif ext == ".doc" and self.include_doc:
                        self.log.emit(f"[DRY-RUN] Would convert+update: {f}")
                    self.progress.emit(idx, len(files))
                    continue

                work_docx, original_doc = None, None
                if ext == ".docx":
                    work_docx = f
                elif ext == ".doc":
                    if not self.include_doc:
                        self.log.emit(f"[SKIP] (legacy .doc; enable 'Include .doc') {f}")
                        self.progress.emit(idx, len(files))
                        continue
                    work_docx = f.with_suffix(".docx")
                    original_doc = f
                    convert_doc_to_docx(word, f, work_docx)

                result = update_docx_dates(work_docx, target_date, self.phase_text,
                                               normalize_fonts=self.normalize_fonts, 
                                               target_font=self.target_font,
                                               target_font_size=self.target_font_size)
                if result['changed']:
                    updated_ct += 1
                    self.stats['documents_updated'] += 1
                    
                    # Record changes
                    if result['date_changed']:
                        self.stats['documents_with_date_changes'] += 1
                    if result['phase_changed']:
                        self.stats['documents_with_phase_changes'] += 1
                    if result.get('fonts_changed'):
                        self.stats['documents_with_font_changes'] += 1
                    if result['date_changed'] and result['phase_changed']:
                        self.stats['documents_with_both'] += 1
                    
                    # Build update message with details
                    changes = []
                    if result['date_changed']: changes.append('date')
                    if result['phase_changed']: changes.append('phase')
                    if result.get('fonts_changed'): changes.append('fonts')
                    change_str = ', '.join(changes) if changes else 'content'
                    self.log.emit(f"[UPDATED: {change_str}] {f}")

                    if original_doc and self.replace_doc_inplace:
                        try:
                            original_doc.unlink(missing_ok=True)
                        except Exception:
                            pass

                    if self.reprint_pdf:
                        pdf_path = work_docx.with_suffix(".pdf")
                        try:
                            if pdf_path.exists():
                                pdf_path.unlink()
                        except Exception:
                            try:
                                pdf_path.rename(pdf_path.with_suffix(".pdf.bak"))
                            except Exception:
                                pass
                        export_pdf(word, work_docx, pdf_path)
                        self.stats['pdfs_created'] += 1
                        self.log.emit(f"  -> [PDF REPRINTED] {pdf_path}")
                else:
                    # Only log "no changes" if we weren't just doing font normalization
                    if not self.normalize_fonts:
                        self.log.emit(f"[NO CHANGES] {f}")
                    else:
                        self.log.emit(f"[NO CHANGES] {f}")

            except Exception as e:
                errors += 1
                self.stats['errors'] += 1
                self.log.emit(f"[ERROR] {f} -> {e}")

            self.progress.emit(idx, len(files))
        
        # Finalize statistics
        self.stats['files_scanned'] = len(files)
        self.stats['errors'] = errors
        if self.start_time:
            self.stats['duration_seconds'] = time.time() - self.start_time

        safe_close_word(word)
        # kill_orphaned_winword()  # keep disabled by default
        self.finished.emit(updated_ct, errors, self.stats)
        self.enableUI.emit(True)


# ----------------------- GUI -----------------------
class SubscriptionDialog(QDialog):
    """Modern dialog for entering and validating subscription key."""
    def __init__(self, parent=None, subscription_mgr=None, required=True):
        super().__init__(parent)
        self.subscription_mgr = subscription_mgr
        self.required = required
        self.validation_result = None
        
        self.setWindowTitle("🔑 License Activation" if required else "🔑 Enter License Key")
        self.setModal(True)
        self.setMinimumWidth(550)
        
        # Apply theme-aware styling
        self._apply_theme()
        
        layout = QVBoxLayout(self)
        layout.setSpacing(15)
        
        # Header with icon
        header_layout = QHBoxLayout()
        icon_label = QLabel("🔑")
        icon_label.setStyleSheet("font-size: 40px;")
        
        title_label = QLabel("License Activation Required" if required else "Activate License")
        title_label.setStyleSheet(f"""
            font-size: 20px;
            font-weight: bold;
            color: {Colors.TEXT};
        """)
        
        header_layout.addWidget(icon_label)
        header_layout.addWidget(title_label, 1)
        layout.addLayout(header_layout)
        
        # Info message
        if required:
            info_text = "This application requires a valid license to run. Please enter your license key below."
        else:
            info_text = "Enter your license key to unlock premium features."
        
        info_label = QLabel(info_text)
        info_label.setWordWrap(True)
        info_label.setStyleSheet(f"""
            color: {Colors.TEXT_SECONDARY};
            padding: 12px;
            background: {Colors.CARD};
            border: 1px solid {Colors.BORDER};
            border-radius: 6px;
            font-size: 13px;
        """)
        layout.addWidget(info_label)
        
        # License key input with validation feedback
        key_group = QGroupBox("License Key")
        key_layout = QVBoxLayout()
        
        self.license_key_input = QLineEdit()
        self.license_key_input.setPlaceholderText("XXXXX-XXXXX-XXXXX-XXXXX")
        self.license_key_input.textChanged.connect(self.on_key_changed)
        self.license_key_input.setStyleSheet(f"""
            QLineEdit {{
                border: 2px solid {Colors.BORDER};
                border-radius: 6px;
                padding: 12px;
                font-size: 15px;
                font-family: 'Consolas', 'Monaco', monospace;
                background-color: {Colors.INPUT_BG};
                color: {Colors.TEXT};
            }}
            QLineEdit:focus {{
                border: 2px solid {Colors.PRIMARY};
            }}
        """)
        key_layout.addWidget(self.license_key_input)
        
        # Validation feedback label
        self.feedback_label = QLabel("")
        self.feedback_label.setWordWrap(True)
        self.feedback_label.setStyleSheet("padding: 5px; min-height: 20px;")
        key_layout.addWidget(self.feedback_label)
        
        key_group.setLayout(key_layout)
        layout.addWidget(key_group)
        
        # Progress indicator (hidden by default)
        self.progress_label = QLabel("⏳ Validating license...")
        self.progress_label.setStyleSheet(f"""
            color: {Colors.TEXT_SECONDARY};
            font-style: italic;
            font-size: 13px;
        """)
        self.progress_label.hide()
        layout.addWidget(self.progress_label)
        
        # Buttons
        button_layout = QHBoxLayout()
        button_layout.setSpacing(10)
        
        self.purchase_btn = ModernButton("🛒 Purchase License", "info")
        self.purchase_btn.clicked.connect(self.show_help)
        
        button_layout.addWidget(self.purchase_btn)
        button_layout.addStretch()
        
        if not required:
            self.cancel_btn = ModernButton("Skip for Now", "secondary")
            self.cancel_btn.clicked.connect(self.reject)
            button_layout.addWidget(self.cancel_btn)
        
        self.activate_btn = ModernButton("✓ Activate License", "success")
        self.activate_btn.clicked.connect(self.validate_and_accept)
        self.activate_btn.setDefault(True)
        self.activate_btn.setEnabled(False)
        button_layout.addWidget(self.activate_btn)
        
        if required:
            self.exit_btn = ModernButton("❌ Exit Application", "danger")
            self.exit_btn.clicked.connect(self.exit_app)
            button_layout.addWidget(self.exit_btn)
        
        layout.addLayout(button_layout)
        
        self.setLayout(layout)
        self.license_key_input.setFocus()
    
    def on_key_changed(self, text):
        """Handle license key text changes."""
        text = text.strip()
        
        # Enable/disable activate button based on input
        has_text = len(text) > 0
        self.activate_btn.setEnabled(has_text)
        
        # Clear previous feedback
        if not text:
            self.feedback_label.setText("")
            self.feedback_label.setStyleSheet("padding: 5px; min-height: 20px;")
        else:
            # Basic format validation (visual feedback only)
            if len(text) >= 5:
                self.feedback_label.setText("✓ Format looks good")
                self.feedback_label.setStyleSheet(f"padding: 5px; color: {Colors.SUCCESS}; min-height: 20px; font-weight: 500;")
            else:
                self.feedback_label.setText("")
                self.feedback_label.setStyleSheet("padding: 5px; min-height: 20px;")
    
    def validate_and_accept(self):
        """Validate the license key before accepting."""
        license_key = self.license_key_input.text().strip()
        
        if not license_key:
            self.show_error("Please enter a license key")
            return
        
        # Show progress
        self.progress_label.show()
        self.activate_btn.setEnabled(False)
        self.license_key_input.setEnabled(False)
        QApplication.processEvents()  # Update UI
        
        # Validate with subscription manager
        if self.subscription_mgr:
            try:
                if self.subscription_mgr.validate_license_key(license_key):
                    self.validation_result = "success"
                    self.show_success("✓ License activated successfully!")
                    # Wait a moment to show success message
                    QTimer.singleShot(800, self.accept)
                else:
                    self.validation_result = "invalid"
                    self.show_error("Invalid or expired license key. Please check and try again.")
                    self.reset_ui()
            except Exception as e:
                self.validation_result = "error"
                self.show_error(f"Validation error: {str(e)}")
                self.reset_ui()
        else:
            # No subscription manager (shouldn't happen)
            self.accept()
    
    def show_error(self, message):
        """Show error feedback."""
        self.feedback_label.setText(f"✗ {message}")
        self.feedback_label.setStyleSheet("padding: 5px; color: #dc3545; font-weight: bold; min-height: 20px;")
        self.progress_label.hide()
    
    def show_success(self, message):
        """Show success feedback."""
        self.feedback_label.setText(message)
        self.feedback_label.setStyleSheet("padding: 5px; color: #28a745; font-weight: bold; min-height: 20px;")
        self.progress_label.hide()
    
    def reset_ui(self):
        """Reset UI after failed validation."""
        self.activate_btn.setEnabled(True)
        self.license_key_input.setEnabled(True)
        self.license_key_input.setFocus()
        self.license_key_input.selectAll()
        self.progress_label.hide()
    
    def closeEvent(self, event):
        """Handle window close event (X button)."""
        # If licensing is required and user hasn't successfully validated, exit the app
        if self.required and self.validation_result != "success":
            # Prevent dialog from closing, exit the application instead
            event.ignore()
            sys.exit(0)
        else:
            # Not required, or validation succeeded - allow normal close
            event.accept()
    
    def exit_app(self):
        """Exit the application."""
        sys.exit(0)
    
    def _apply_theme(self):
        """Apply theme-aware styling to dialog."""
        self.setStyleSheet(f"""
            QDialog {{
                background-color: {Colors.BACKGROUND};
            }}
            QLabel {{
                color: {Colors.TEXT};
            }}
            QGroupBox {{
                border: 2px solid {Colors.BORDER};
                border-radius: 8px;
                margin-top: 10px;
                padding-top: 15px;
                font-weight: bold;
                background-color: {Colors.CARD};
                color: {Colors.TEXT};
            }}
            QGroupBox::title {{
                color: {Colors.TEXT};
                subcontrol-origin: margin;
                left: 15px;
                padding: 0 5px;
            }}
            QLineEdit {{
                background-color: {Colors.INPUT_BG};
                color: {Colors.TEXT};
                border: 2px solid {Colors.BORDER};
                border-radius: 6px;
                padding: 12px;
            }}
            QLineEdit:focus {{
                border: 2px solid {Colors.PRIMARY};
            }}
        """)
    
    def show_help(self):
        """Open purchase URL in default browser."""
        webbrowser.open("https://example.com/purchase-license")


class SubscriptionStatusWidget(QFrame):
    """Modern widget to display subscription status."""
    def __init__(self, subscription_mgr, parent=None):
        super().__init__(parent)
        self.subscription_mgr = subscription_mgr
        
        # Modern card styling
        self.setStyleSheet(f"""
            QFrame {{
                background-color: {Colors.CARD};
                border: 1px solid {Colors.BORDER};
                border-radius: 5px;
                padding: 6px 8px;
            }}
        """)
        
        layout = QHBoxLayout()
        layout.setContentsMargins(0, 0, 0, 0)
        layout.setSpacing(8)
        
        self.status_icon = QLabel("🔒")
        self.status_icon.setStyleSheet("font-size: 16px;")
        
        self.status_label = QLabel("Checking subscription...")
        self.status_label.setStyleSheet(f"""
            font-size: 11px;
            color: {Colors.TEXT};
            font-weight: 500;
        """)
        
        self.manage_btn = ModernButton("⚙️", "secondary")
        self.manage_btn.setToolTip("Manage subscription")
        self.manage_btn.setFixedWidth(28)
        self.manage_btn.clicked.connect(self.manage_subscription)
        
        layout.addWidget(self.status_icon)
        layout.addWidget(self.status_label, 1)
        layout.addWidget(self.manage_btn)
        
        self.setLayout(layout)
        self.update_style()
        self.update_status()
    
    def update_style(self):
        """Update widget styling for current theme."""
        self.setStyleSheet(f"""
            QFrame {{
                background-color: {Colors.CARD};
                border: 1px solid {Colors.BORDER};
                border-radius: 5px;
                padding: 6px 8px;
            }}
        """)
        self.status_label.setStyleSheet(f"""
            font-size: 11px;
            color: {Colors.TEXT};
            font-weight: 500;
        """)
        if hasattr(self, 'manage_btn'):
            self.manage_btn.update_style()
    
    def update_status(self):
        """Update the status display."""
        info = self.subscription_mgr.get_subscription_info()
        
        if info['status'] == 'active':
            self.status_icon.setText("✅")
            expiry_str = info.get('expiry_date')
            plan = info.get('plan', 'unknown')
            
            if expiry_str and expiry_str != 'None':
                try:
                    expiry = datetime.fromisoformat(expiry_str).strftime("%Y-%m-%d")
                    expiry_text = f"Active until {expiry}"
                except (ValueError, TypeError):
                    expiry_text = "Active"
            else:
                # Free license or no expiration
                if plan == 'free':
                    expiry_text = "Free License (Unlimited)"
                else:
                    expiry_text = "Active (No expiration)"
            
            if info['documents_remaining'] is None or info['documents_remaining'] < 0:
                self.status_label.setText(f"{expiry_text} - Unlimited documents")
            else:
                self.status_label.setText(
                    f"{expiry_text} - {info['documents_remaining']} docs remaining"
                )
        else:
            self.status_icon.setText("⚠️")
            self.status_label.setText("No active subscription")
    
    def manage_subscription(self):
        """Open subscription management in browser."""
        webbrowser.open("https://example.com/manage-subscription")


class MainWindow(QWidget):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("📝 Spec Header Date & Phase Updater")
        self.resize(900, 500)
        
        # Load app configuration
        self.app_config = self._load_app_config()
        
        # Load and apply theme FIRST before any styling
        dark_mode = self.app_config.get('dark_mode', False)
        Colors.set_theme(dark_mode)
        
        # Set modern background AFTER theme is applied
        self.setStyleSheet(f"""
            QWidget {{
                background-color: {Colors.BACKGROUND};
                color: {Colors.TEXT};
                font-size: 12px;
            }}
            QLabel {{
                color: {Colors.TEXT};
                font-weight: 500;
                font-size: 11px;
            }}
        """)
        
        # Initialize subscription manager (conditional)
        if INCLUDE_LICENSING:
            self.subscription_mgr = SubscriptionManager()
        else:
            # Dummy manager when licensing disabled
            self.subscription_mgr = SubscriptionManager()

        # Apply modern input styling helper
        self._setup_modern_styles()
        
        # Inputs
        self.txtRoot = QLineEdit()
        self.txtRoot.setStyleSheet(self._input_style())
        self.txtRoot.setMinimumHeight(28)
        self.txtRoot.setPlaceholderText("Select specifications folder...")
        
        self.btnBrowseRoot = ModernButton("📁", "secondary")
        self.btnBrowseRoot.setToolTip("Browse for folder")
        self.btnBrowseRoot.setFixedWidth(32)
        
        self.dateEdit = QDateEdit()
        self.dateEdit.setCalendarPopup(True)
        self.dateEdit.setDisplayFormat("MMMM d, yyyy")
        self.dateEdit.setDate(QDate.currentDate())
        self.dateEdit.setStyleSheet(self._input_style())
        self.dateEdit.setMinimumHeight(28)
        
        self.txtPhase = QLineEdit()
        self.txtPhase.setText(DEFAULT_PHASE_TEXT)
        self.txtPhase.setPlaceholderText("e.g., 100% Construction Documents")
        self.txtPhase.setStyleSheet(self._input_style())
        self.txtPhase.setMinimumHeight(28)

        self.chkRecursive = QCheckBox("Include files in subfolders")
        self.chkRecursive.setChecked(True)
        self.chkRecursive.setStyleSheet(self._checkbox_style())

        self.chkIncludeDoc = QCheckBox("Include legacy .doc (requires Word)")
        self.chkReplaceDoc = QCheckBox("Replace .doc with updated .docx (delete original)")
        self.chkReprintPDF = QCheckBox("Reprint same-name PDFs (requires Word)")
        
        self.chkReprintOnly = QCheckBox("Reprint PDFs only (no doc changes)")
        self.chkReprintOnly.setChecked(False)
        self.chkReprintOnly.setToolTip("Only regenerate PDFs from Word documents, no edits applied")

        self.chkDryRun = QCheckBox("Dry-run (no edits)")
        self.chkDryRun.setChecked(False)
        
        self.chkSkipTOC = QCheckBox("Skip 'Table of Contents'")
        self.chkSkipTOC.setChecked(True)
        self.chkSkipTOC.setToolTip("Skip files with 'Table of Contents' in the filename")
        
        # Font normalization
        self.chkNormalizeFonts = QCheckBox("Normalize fonts")
        self.chkNormalizeFonts.setChecked(False)
        self.chkNormalizeFonts.setToolTip("Standardize all fonts to a single font")
        
        self.txtTargetFont = QLineEdit()
        self.txtTargetFont.setText(DEFAULT_FONT_NAME)
        self.txtTargetFont.setPlaceholderText("e.g., Arial, Times New Roman")
        self.txtTargetFont.setStyleSheet(self._input_style())
        self.txtTargetFont.setMinimumHeight(28)
        self.txtTargetFont.setMaximumWidth(120)
        self.txtTargetFont.setEnabled(False)  # Disabled until checkbox is checked
        
        # Font size normalization
        self.chkNormalizeFontSize = QCheckBox("Set uniform size")
        self.chkNormalizeFontSize.setChecked(False)
        self.chkNormalizeFontSize.setToolTip("Also set all text to the same font size")
        self.chkNormalizeFontSize.setEnabled(False)
        
        self.spnFontSize = QSpinBox()
        self.spnFontSize.setRange(6, 72)
        self.spnFontSize.setValue(DEFAULT_FONT_SIZE)
        self.spnFontSize.setSuffix(" pt")
        self.spnFontSize.setMinimumHeight(28)
        self.spnFontSize.setMaximumWidth(70)
        self.spnFontSize.setEnabled(False)
        self.spnFontSize.setStyleSheet(self._input_style())
        
        # Connect font normalization checkbox to enable/disable related controls
        def on_normalize_fonts_toggled(checked):
            self.txtTargetFont.setEnabled(checked)
            self.chkNormalizeFontSize.setEnabled(checked)
            if not checked:
                self.chkNormalizeFontSize.setChecked(False)
                self.spnFontSize.setEnabled(False)
        
        self.chkNormalizeFonts.toggled.connect(on_normalize_fonts_toggled)
        self.chkNormalizeFontSize.toggled.connect(self.spnFontSize.setEnabled)

        # Backup
        self.txtBackup = QLineEdit()
        self.txtBackup.setStyleSheet(self._input_style())
        self.txtBackup.setMinimumHeight(28)
        self.txtBackup.setPlaceholderText("Select backup folder...")
        
        self.btnBrowseBackup = ModernButton("📁", "secondary")
        self.btnBrowseBackup.setToolTip("Browse for backup folder")
        self.btnBrowseBackup.setFixedWidth(32)
        
        self.chkUseBackup = QCheckBox("Save backups before editing")
        self.chkUseBackup.setChecked(False)
        self.chkUseBackup.setStyleSheet(self._checkbox_style())

        # Exclude folders
        self.lstExclude = QListWidget()
        self.lstExclude.setStyleSheet(f"""
            QListWidget {{
                background-color: {Colors.CARD};
                border: 1px solid {Colors.BORDER};
                border-radius: 4px;
                padding: 4px;
                font-size: 11px;
            }}
            QListWidget::item {{
                padding: 4px;
                border-radius: 2px;
            }}
            QListWidget::item:selected {{
                background-color: {Colors.PRIMARY};
                color: white;
            }}
        """)
        self.lstExclude.setMaximumHeight(70)
        for name in ("_archive", "archive"):
            self.lstExclude.addItem(QListWidgetItem(name))
        self.btnAddEx = ModernButton("➕ Add", "secondary")
        self.btnRemoveEx = ModernButton("➖ Remove", "secondary")

        # Controls
        self.btnRun = ModernButton("▶️ Run Update", "success")
        self.btnCancel = ModernButton("⏹️ Cancel", "danger")
        self.btnCancel.setEnabled(False)

        self.progress = QProgressBar()
        self.progress.setMinimum(0)
        self.progress.setMaximum(100)
        self.progress.setValue(0)
        self.progress.setStyleSheet(f"""
            QProgressBar {{
                border: 1px solid {Colors.BORDER};
                border-radius: 4px;
                text-align: center;
                background-color: {Colors.CARD};
                height: 26px;
                font-weight: 600;
                font-size: 10px;
            }}
            QProgressBar::chunk {{
                background-color: {Colors.SUCCESS};
                border-radius: 2px;
            }}
        """)

        self.log = QTextEdit()
        self.log.setReadOnly(True)
        self.log.setMaximumHeight(80)
        self.log.setStyleSheet(f"""
            QTextEdit {{
                background-color: {Colors.CARD};
                border: 1px solid {Colors.BORDER};
                border-radius: 4px;
                padding: 5px;
                font-family: 'Consolas', 'Monaco', monospace;
                font-size: 10px;
                line-height: 1.35;
            }}
        """)

        # Subscription status bar (conditional, hidden by default)
        if INCLUDE_LICENSING:
            self.subscription_status = SubscriptionStatusWidget(self.subscription_mgr)
            self.subscription_status.setVisible(False)  # Hidden by default
        else:
            self.subscription_status = None  # No status widget when licensing disabled
        
        # Layout - Compact two-column design
        main_layout = QVBoxLayout()
        main_layout.setContentsMargins(8, 8, 8, 8)
        main_layout.setSpacing(5)
        
        # Top bar with subscription status and dark mode toggle - aligned to the right
        top_bar_widget = QWidget()
        top_bar_widget.setFixedHeight(32)
        top_bar_widget.setStyleSheet("background-color: transparent;")
        top_bar = QHBoxLayout(top_bar_widget)
        top_bar.setContentsMargins(0, 0, 0, 0)
        top_bar.setSpacing(8)
        
        # Add subscription status bar at the top (hidden by default)
        if self.subscription_status:
            top_bar.addWidget(self.subscription_status, 1)
        
        # Add stretch to push dark mode button to the right
        top_bar.addStretch()
        
        # Dark mode toggle button in top right
        self.btnDarkMode = ModernButton("🌙" if not dark_mode else "☀️", "secondary")
        self.btnDarkMode.setToolTip("Toggle dark mode" if not dark_mode else "Toggle light mode")
        self.btnDarkMode.setFixedWidth(32)
        self.btnDarkMode.setFixedHeight(26)
        self.btnDarkMode.setCheckable(True)
        self.btnDarkMode.setChecked(dark_mode)
        self.btnDarkMode.clicked.connect(self.toggle_dark_mode)
        # Force update style after theme is set to ensure correct colors
        self.btnDarkMode.update_style()
        top_bar.addWidget(self.btnDarkMode)
        
        main_layout.addWidget(top_bar_widget)
        
        # Main Settings Card
        settings_card = ModernCard("⚙️ Update Settings")
        settings_grid = QGridLayout()
        settings_grid.setSpacing(5)
        settings_grid.setVerticalSpacing(4)
        
        settings_grid.addWidget(QLabel("Specs Folder:"), 0, 0, Qt.AlignRight)
        settings_grid.addWidget(self.txtRoot, 0, 1)
        settings_grid.addWidget(self.btnBrowseRoot, 0, 2)
        
        settings_grid.addWidget(QLabel("New Date:"), 1, 0, Qt.AlignRight)
        settings_grid.addWidget(self.dateEdit, 1, 1)
        
        settings_grid.addWidget(QLabel("Phase Text:"), 2, 0, Qt.AlignRight)
        settings_grid.addWidget(self.txtPhase, 2, 1)
        # Inline recursive option on the same row to save vertical space
        settings_grid.addWidget(self.chkRecursive, 2, 2)
        
        settings_card.content_layout.addLayout(settings_grid)
        main_layout.addWidget(settings_card)
        
        # Options in compact inline format
        options_card = ModernCard("🔧 Options")
        options_grid = QGridLayout()
        options_grid.setSpacing(3)
        options_grid.setVerticalSpacing(4)
        for i, chk in enumerate([self.chkIncludeDoc, self.chkReplaceDoc, self.chkReprintPDF, self.chkReprintOnly, self.chkSkipTOC]):
            chk.setStyleSheet(self._checkbox_style())
            options_grid.addWidget(chk, 0, i)
        
        # Second row: dry run and font normalization
        self.chkDryRun.setStyleSheet(self._checkbox_style())
        self.chkNormalizeFonts.setStyleSheet(self._checkbox_style())
        self.chkNormalizeFontSize.setStyleSheet(self._checkbox_style())
        options_grid.addWidget(self.chkDryRun, 1, 0)
        options_grid.addWidget(self.chkNormalizeFonts, 1, 1)
        
        font_row = QHBoxLayout()
        font_row.setSpacing(6)
        font_label = QLabel("Font:")
        font_label.setStyleSheet(f"color: {Colors.TEXT_SECONDARY}; font-size: 10px;")
        font_row.addWidget(font_label)
        font_row.addWidget(self.txtTargetFont)
        font_row.addSpacing(10)
        font_row.addWidget(self.chkNormalizeFontSize)
        font_row.addWidget(self.spnFontSize)
        font_row.addStretch()
        options_grid.addLayout(font_row, 1, 2, 1, 3)
        
        options_card.content_layout.addLayout(options_grid)
        main_layout.addWidget(options_card)
        
        # Two-column layout for Backup and Exclude
        bottom_row = QHBoxLayout()
        bottom_row.setSpacing(4)
        
        # Backup Card (left)
        backup_card = ModernCard("💾 Backup")
        backup_grid = QGridLayout()
        backup_grid.setSpacing(6)
        backup_grid.addWidget(self.chkUseBackup, 0, 0, 1, 3)
        backup_grid.addWidget(QLabel("Folder:"), 1, 0, Qt.AlignRight)
        backup_grid.addWidget(self.txtBackup, 1, 1)
        backup_grid.addWidget(self.btnBrowseBackup, 1, 2)
        backup_card.content_layout.addLayout(backup_grid)
        bottom_row.addWidget(backup_card)
        
        # Exclude Folders Card (right)
        exclude_card = ModernCard("🚫 Exclude")
        exclude_grid = QGridLayout()
        exclude_grid.setSpacing(4)
        exclude_grid.addWidget(self.lstExclude, 0, 0, 3, 1)
        exclude_grid.addWidget(self.btnAddEx, 0, 1)
        exclude_grid.addWidget(self.btnRemoveEx, 1, 1)
        exclude_card.content_layout.addLayout(exclude_grid)
        bottom_row.addWidget(exclude_card)
        
        main_layout.addLayout(bottom_row)
        
        # Progress + Actions in one row
        control_row = QHBoxLayout()
        control_row.setSpacing(6)
        control_row.addWidget(self.progress, 1)
        control_row.addWidget(self.btnRun)
        control_row.addWidget(self.btnCancel)
        main_layout.addLayout(control_row)
        
        # Log (no card wrapper to save space)
        log_label = QLabel("📋 Log")
        log_label.setStyleSheet(f"font-weight: bold; font-size: 12px; color: {Colors.TEXT}; padding: 3px 0;")
        main_layout.addWidget(log_label)
        main_layout.addWidget(self.log)
        
        self.setLayout(main_layout)

        # Connections
        self.btnBrowseRoot.clicked.connect(self.pickRoot)
        self.btnBrowseBackup.clicked.connect(self.pickBackup)
        self.btnAddEx.clicked.connect(self.addExclude)
        self.btnRemoveEx.clicked.connect(self.removeExclude)
        self.btnRun.clicked.connect(self.startRun)
        self.btnCancel.clicked.connect(self.cancelRun)
        
        self.worker: UpdateWorker | None = None
        
        # Check subscription on startup (after UI is fully created) - conditional
        if INCLUDE_LICENSING:
            self.check_subscription()
            
            # Set up a timer to periodically check subscription status
            self.subscription_timer = QTimer(self)
            self.subscription_timer.timeout.connect(self.check_subscription)
            self.subscription_timer.start(5 * 60 * 1000)  # Check every 5 minutes
        else:
            self.subscription_timer = None
    
    def _setup_modern_styles(self):
        """Initialize modern styling - placeholder for future enhancements."""
        pass
    
    def _input_style(self) -> str:
        """Return modern input field styling."""
        return f"""
            QLineEdit, QDateEdit {{
                border: 1px solid {Colors.BORDER};
                border-radius: 4px;
                padding: 4px 8px;
                font-size: 12px;
                background-color: {Colors.INPUT_BG};
                color: {Colors.TEXT};
                selection-background-color: {Colors.PRIMARY};
            }}
            QLineEdit:hover, QDateEdit:hover {{
                border-color: {Colors.TEXT_MUTED};
            }}
            QLineEdit:focus, QDateEdit:focus {{
                border: 1px solid {Colors.PRIMARY};
                outline: none;
            }}
            QDateEdit::drop-down {{
                border: none;
                width: 20px;
            }}
            QDateEdit::down-arrow {{
                image: none;
                border-left: 3px solid transparent;
                border-right: 3px solid transparent;
                border-top: 4px solid {Colors.TEXT_SECONDARY};
                margin-right: 4px;
            }}
        """
    
    def _checkbox_style(self) -> str:
        """Return modern checkbox styling."""
        return f"""
            QCheckBox {{
                color: {Colors.TEXT};
                font-size: 11px;
                spacing: 4px;
                padding: 2px;
                background-color: transparent;
                border: none;
            }}
            QCheckBox:hover {{
                color: {Colors.PRIMARY};
                background-color: transparent;
            }}
            QCheckBox::indicator {{
                width: 14px;
                height: 14px;
                border: 1px solid {Colors.BORDER};
                border-radius: 3px;
                background-color: {Colors.CARD};
            }}
            QCheckBox::indicator:hover {{
                border-color: {Colors.PRIMARY};
                border-radius: 3px;
            }}
            QCheckBox::indicator:checked {{
                background-color: {Colors.PRIMARY};
                border-color: {Colors.PRIMARY};
                border-radius: 3px;
                image: none;
            }}
        """

    def _center_dialog(self, dialog):
        """Center a dialog on the main window."""
        if dialog and self.isVisible():
            parent_geo = self.geometry()
            dialog_size = dialog.sizeHint()
            x = parent_geo.x() + (parent_geo.width() - dialog_size.width()) // 2
            y = parent_geo.y() + (parent_geo.height() - dialog_size.height()) // 2
            dialog.move(x, y)
    
    def _load_app_config(self) -> dict:
        """Load application configuration."""
        # Check multiple locations for app_config.json
        config_locations = [
            Path(__file__).parent.parent / 'app_config.json',  # Renamer root folder
            Path(__file__).parent / 'app_config.json',  # src folder (fallback)
            Path(__file__).parent.parent.parent / 'app_config.json',  # desktop-widgets folder
        ]
        
        config_file = None
        for path in config_locations:
            if path.exists():
                config_file = path
                break
        
        default_config = {
            "require_subscription": True,
            "app_name": "Spec Header Date Updater",
            "app_version": "1.0.0",
            "dark_mode": False,
            "min_plan": "free"
        }
        
        try:
            if config_file and config_file.exists():
                with open(config_file, 'r') as f:
                    return json.load(f)
        except Exception as e:
            print(f"Error loading config: {e}")
        
        return default_config
    
    def _save_app_config(self) -> None:
        """Save application configuration."""
        config_locations = [
            Path(__file__).parent.parent / 'app_config.json',
            Path(__file__).parent / 'app_config.json',
            Path(__file__).parent.parent.parent / 'app_config.json',
        ]
        
        config_file = None
        for path in config_locations:
            if path.exists():
                config_file = path
                break
        
        if not config_file:
            config_file = Path(__file__).parent.parent / 'app_config.json'
        
        try:
            with open(config_file, 'w') as f:
                json.dump(self.app_config, f, indent=2)
        except Exception as e:
            print(f"Error saving config: {e}")
    
    def toggle_dark_mode(self):
        """Toggle dark mode on/off."""
        dark_mode = self.btnDarkMode.isChecked()
        Colors.set_theme(dark_mode)
        self.app_config['dark_mode'] = dark_mode
        self._save_app_config()
        self._refresh_all_styles()
    
    def _refresh_all_styles(self):
        """Refresh all widget styles after theme change."""
        # Update main window style
        self.setStyleSheet(f"""
            QWidget {{
                background-color: {Colors.BACKGROUND};
                color: {Colors.TEXT};
                font-size: 12px;
            }}
            QLabel {{
                color: {Colors.TEXT};
                font-weight: 500;
                font-size: 11px;
            }}
            QMessageBox {{
                background-color: {Colors.BACKGROUND};
                color: {Colors.TEXT};
            }}
            QMessageBox QLabel {{
                color: {Colors.TEXT};
            }}
            QCalendarWidget {{
                background-color: {Colors.CARD};
                color: {Colors.TEXT};
            }}
            QCalendarWidget QTableView {{
                selection-background-color: {Colors.PRIMARY};
                selection-color: white;
            }}
        """)
        
        # Update all input fields
        for widget in [self.txtRoot, self.txtPhase, self.txtBackup, self.dateEdit]:
            widget.setStyleSheet(self._input_style())
        
        # Update all checkboxes
        for checkbox in [self.chkRecursive, self.chkIncludeDoc, self.chkReplaceDoc, 
                        self.chkReprintPDF, self.chkDryRun, self.chkUseBackup]:
            checkbox.setStyleSheet(self._checkbox_style())
        
        # Update list widget
        self.lstExclude.setStyleSheet(f"""
            QListWidget {{
                background-color: {Colors.CARD};
                border: 1px solid {Colors.BORDER};
                border-radius: 4px;
                padding: 4px;
                font-size: 11px;
                color: {Colors.TEXT};
            }}
            QListWidget::item {{
                padding: 4px;
                border-radius: 2px;
                color: {Colors.TEXT};
            }}
            QListWidget::item:selected {{
                background-color: {Colors.PRIMARY};
                color: white;
            }}
            QListWidget::item:hover {{
                background-color: {Colors.BORDER};
            }}
        """)
        
        # Update progress bar
        self.progress.setStyleSheet(f"""
            QProgressBar {{
                border: 1px solid {Colors.BORDER};
                border-radius: 4px;
                text-align: center;
                background-color: {Colors.CARD};
                height: 26px;
                font-weight: 600;
                font-size: 10px;
                color: {Colors.TEXT};
            }}
            QProgressBar::chunk {{
                background-color: {Colors.SUCCESS};
                border-radius: 2px;
            }}
        """)
        
        # Update log
        self.log.setStyleSheet(f"""
            QTextEdit {{
                background-color: {Colors.CARD};
                border: 1px solid {Colors.BORDER};
                border-radius: 4px;
                padding: 5px;
                font-family: 'Consolas', 'Monaco', monospace;
                font-size: 10px;
                line-height: 1.35;
                color: {Colors.TEXT};
            }}
        """)
        
        # Update log label
        log_labels = self.findChildren(QLabel)
        for label in log_labels:
            if label.text() == "📋 Log":
                label.setStyleSheet(f"font-weight: bold; font-size: 12px; color: {Colors.TEXT}; padding: 3px 0;")
        
        # Refresh all ModernCard widgets
        for widget in self.findChildren(ModernCard):
            widget.update_style()
        
        # Refresh all ModernButton widgets
        for widget in self.findChildren(ModernButton):
            widget.update_style()
        
        # Refresh subscription status widget
        if hasattr(self, 'subscription_status') and self.subscription_status:
            self.subscription_status.update_style()
        
        # Update button icon based on theme
        dark_mode = self.btnDarkMode.isChecked()
        if dark_mode:
            self.btnDarkMode.setText("☀️")
            self.btnDarkMode.setToolTip("Toggle light mode")
        else:
            self.btnDarkMode.setText("🌙")
            self.btnDarkMode.setToolTip("Toggle dark mode")
    
    def check_subscription(self):
        """Check subscription status and prompt for license if needed."""
        # Skip all subscription checks if licensing is disabled
        if not INCLUDE_LICENSING:
            return
        
        require_sub = self.app_config.get('require_subscription', True)
        
        # Always ensure license exists (auto-create free license if needed)
        if not self.subscription_mgr.is_subscribed():
            # Try to auto-create free license first (silent)
            if not self.subscription_mgr.ensure_license_exists():
                # If auto-create failed and subscription is required, show dialog
                if require_sub:
                    dialog = SubscriptionDialog(
                        parent=self,
                        subscription_mgr=self.subscription_mgr,
                        required=True
                    )
                    dialog.exec()
        
        # Sync activation status with license server
        try:
            self.subscription_mgr._sync_activation_status()
        except Exception:
            pass  # Silent fail
        
        # Update the UI
        self.update_subscription_ui()
        
        # Disable features if subscription required but not active
        if INCLUDE_LICENSING:
            require_sub = self.app_config.get('require_subscription', True)
            if require_sub and not self.subscription_mgr.is_subscribed():
                self.btnRun.setEnabled(False)
                self.btnRun.setToolTip("Valid subscription required to process documents")
    
    def update_subscription_ui(self):
        """Update UI elements based on subscription status."""
        # Skip if licensing disabled
        if not INCLUDE_LICENSING or not self.subscription_status:
            return
        
        # Update the status widget
        self.subscription_status.update_status()
        
        # Optionally disable features for non-subscribers
        # For now, we'll allow basic functionality even without subscription
        # You can uncomment below to require subscription for all features:
        # self.btnRun.setEnabled(self.subscription_mgr.is_subscribed())

    @Slot()
    def pickRoot(self):
        path = QFileDialog.getExistingDirectory(self, "Select Specifications Folder")
        if path:
            self.txtRoot.setText(path)

    @Slot()
    def pickBackup(self):
        path = QFileDialog.getExistingDirectory(self, "Select Backup Folder")
        if path:
            self.txtBackup.setText(path)

    @Slot()
    def addExclude(self):
        path = QFileDialog.getExistingDirectory(self, "Pick a folder name to extract (just name will be used)")
        if path:
            name = Path(path).name
            self.lstExclude.addItem(QListWidgetItem(name))

    @Slot()
    def removeExclude(self):
        for item in self.lstExclude.selectedItems():
            self.lstExclude.takeItem(self.lstExclude.row(item))

    def currentExcludeList(self) -> list[str]:
        return [self.lstExclude.item(i).text() for i in range(self.lstExclude.count())]

    def setUIEnabled(self, en: bool):
        for w in [
            self.txtRoot, self.btnBrowseRoot, self.dateEdit, self.chkRecursive,
            self.chkIncludeDoc, self.chkReplaceDoc, self.chkReprintPDF, self.chkDryRun,
            self.txtBackup, self.btnBrowseBackup, self.chkUseBackup,
            self.lstExclude, self.btnAddEx, self.btnRemoveEx, self.btnRun
        ]:
            w.setEnabled(en)
        self.btnCancel.setEnabled(not en)

    def appendLog(self, msg: str):
        self.log.append(msg)
        from PySide6.QtGui import QTextCursor
        self.log.moveCursor(QTextCursor.MoveOperation.End)

    def _scan_for_legacy_doc_files(self, root: Path, recursive: bool, exclude_folders: list[str], skip_toc: bool) -> int:
        """Quick scan to count legacy .doc files in the target directory."""
        count = 0
        iterator = root.rglob("*.doc") if recursive else root.glob("*.doc")
        for f in iterator:
            if f.name.startswith("~$"):
                continue
            if any(part.lower() in exclude_folders for part in f.parts):
                continue
            if skip_toc and "table of contents" in f.stem.lower():
                continue
            count += 1
        return count

    @Slot()
    def startRun(self):
        root = self.txtRoot.text().strip()
        if not root:
            QMessageBox.warning(self, "Missing folder", "Please choose the specifications folder.")
            return

        date_str = self.dateEdit.date().toString("MMMM d, yyyy")
        recursive = self.chkRecursive.isChecked()
        include_doc = self.chkIncludeDoc.isChecked()
        replace_doc = self.chkReplaceDoc.isChecked()
        reprint_pdf = self.chkReprintPDF.isChecked()
        dry_run = self.chkDryRun.isChecked()
        backup_dir = self.txtBackup.text().strip() if self.chkUseBackup.isChecked() else None
        exclude = self.currentExcludeList()

        self.log.clear()
        self.progress.setValue(0)

        phase_text = self.txtPhase.text().strip() or DEFAULT_PHASE_TEXT
        
        # Check limits before starting worker (if licensing enabled)
        subscription_mgr_for_worker = None
        if INCLUDE_LICENSING and hasattr(self, 'subscription_mgr'):
            subscription_mgr_for_worker = self.subscription_mgr
            
            # Ensure license exists (auto-create if needed)
            if not self.subscription_mgr.is_subscribed():
                self.subscription_mgr.ensure_license_exists()
        
        # Get default backup directory for comparison
        default_backup_dir = self.txtBackup.text().strip() if hasattr(self, 'txtBackup') else None
        
        # Font normalization options
        normalize_fonts = self.chkNormalizeFonts.isChecked()
        target_font = self.txtTargetFont.text().strip() or DEFAULT_FONT_NAME
        target_font_size = self.spnFontSize.value() if self.chkNormalizeFontSize.isChecked() else None
        skip_toc = self.chkSkipTOC.isChecked()
        reprint_only = self.chkReprintOnly.isChecked()
        
        # Pre-scan for legacy .doc files if not included
        if not include_doc:
            root_path = Path(root)
            exclude_lower = [x.lower() for x in exclude]
            doc_count = self._scan_for_legacy_doc_files(root_path, recursive, exclude_lower, skip_toc)
            
            if doc_count > 0:
                # Check if Word is available (unless admin disabled the check)
                require_word_check = self.app_config.get('require_word_check', True)
                word_available, word_reason = check_word_available()
                
                if require_word_check and not word_available:
                    # Word not available - inform user
                    mb = QMessageBox(self)
                    mb.setIcon(QMessageBox.Icon.Warning)
                    mb.setWindowTitle("Legacy Files Detected - Word Required")
                    mb.setText(f"Found {doc_count} legacy .doc file(s) in the target directory.")
                    mb.setInformativeText(
                        f"Legacy .doc files require Microsoft Word to process.\n\n"
                        f"Issue: {word_reason}\n\n"
                        f"Options:\n"
                        f"• Install Microsoft Word to process .doc files\n"
                        f"• Convert .doc files to .docx manually\n"
                        f"• Skip the {doc_count} .doc file(s) and process only .docx files"
                    )
                    btn_skip = mb.addButton(f"Skip .doc Files", QMessageBox.ButtonRole.AcceptRole)
                    btn_cancel = mb.addButton("Cancel", QMessageBox.ButtonRole.RejectRole)
                    mb.setDefaultButton(btn_skip)
                    
                    self._center_dialog(mb)
                    result = mb.exec()
                    clicked = mb.clickedButton()
                    
                    if clicked == btn_cancel or clicked is None:
                        return
                else:
                    # Word available or check disabled - show normal options
                    mb = QMessageBox(self)
                    mb.setIcon(QMessageBox.Icon.Warning)
                    mb.setWindowTitle("Legacy Files Detected")
                    mb.setText(f"Found {doc_count} legacy .doc file(s) in the target directory.")
                    mb.setInformativeText(
                        "Legacy .doc files require Microsoft Word to process.\n\n"
                        "Choose how to handle these files:"
                    )
                    
                    btn_replace = mb.addButton("Convert to .docx (Delete .doc)", QMessageBox.ButtonRole.AcceptRole)
                    btn_keep = mb.addButton("Convert to .docx (Keep .doc)", QMessageBox.ButtonRole.AcceptRole)
                    btn_skip = mb.addButton(f"Skip All .doc Files", QMessageBox.ButtonRole.DestructiveRole)
                    btn_cancel = mb.addButton("Cancel", QMessageBox.ButtonRole.RejectRole)
                    mb.setDefaultButton(btn_replace)
                    
                    self._center_dialog(mb)
                    result = mb.exec()
                    clicked = mb.clickedButton()
                    
                    # Handle X button or Cancel button as cancel
                    if clicked == btn_cancel or clicked is None:
                        return
                    elif clicked == btn_skip:
                        pass
                    elif clicked == btn_replace:
                        include_doc = True
                        replace_doc = True
                        self.chkIncludeDoc.setChecked(True)
                        self.chkReplaceDoc.setChecked(True)
                    elif clicked == btn_keep:
                        include_doc = True
                        replace_doc = False
                        self.chkIncludeDoc.setChecked(True)
                        self.chkReplaceDoc.setChecked(False)
        
        self.worker = UpdateWorker(
            root=root,
            date_str=date_str,
            phase_text=phase_text,
            recursive=recursive,
            dry_run=dry_run,
            backup_dir=backup_dir,
            include_doc=include_doc,
            replace_doc_inplace=replace_doc,
            reprint_pdf=reprint_pdf,
            exclude_folders=exclude,
            subscription_mgr=subscription_mgr_for_worker,
            default_backup_dir=default_backup_dir,
            normalize_fonts=normalize_fonts,
            target_font=target_font,
            target_font_size=target_font_size,
            skip_toc=skip_toc,
            reprint_only=reprint_only
        )
        self.worker.log.connect(self.appendLog)
        self.worker.progress.connect(self.onProgress)
        self.worker.finished.connect(self.onFinished)
        self.worker.needsWord.connect(self.onNeedsWord)
        self.worker.enableUI.connect(self.setUIEnabled)
        self.worker.start()

    @Slot()
    def cancelRun(self):
        if self.worker and self.worker.isRunning():
            self.worker.cancel()
            self.appendLog("[INFO] Cancel requested…")

    @Slot(int, int)
    def onProgress(self, current, total):
        self.progress.setMaximum(total)
        self.progress.setValue(current)

    @Slot(int, int, dict)
    def onFinished(self, updated, errors, stats):
        self.appendLog(f"\nDone. Updated: {updated}, Errors: {errors}")
        
        # Update license usage count (if licensing enabled)
        if INCLUDE_LICENSING and updated > 0 and hasattr(self, 'subscription_mgr') and self.subscription_mgr:
            try:
                self.subscription_mgr.record_document_processed(count=updated)
                # Refresh subscription UI to show updated counts
                self.update_subscription_ui()
            except Exception:
                pass  # Silent fail
        
        # Sync usage statistics with license server
        if INCLUDE_LICENSING and hasattr(self, 'subscription_mgr') and self.subscription_mgr:
            try:
                # Get settings from worker
                worker = self.worker
                default_backup = worker.default_backup_dir
                backup_path = str(worker.backup_dir) if worker.backup_dir else None
                backup_enabled = worker.backup_dir is not None
                backup_location_default = (default_backup and worker.backup_dir and 
                                          str(Path(worker.backup_dir).resolve()) == str(Path(default_backup).resolve()))
                
                # Get exclude folders
                exclude_final = worker.exclude_folders
                
                usage_data = {
                    **stats,  # files_scanned, documents_updated, etc.
                    'date_updated': bool(worker.date_str and worker.date_str.strip()),
                    'phase_updated': bool(worker.phase_text and worker.phase_text.strip()),
                    'target_date': worker.date_str,
                    'phase_text': worker.phase_text,
                    'recursive': worker.recursive,
                    'dry_run': worker.dry_run,
                    'include_legacy_doc': worker.include_doc,
                    'replace_doc_inplace': worker.replace_doc_inplace,
                    'reprint_pdf': worker.reprint_pdf,
                    'backup_enabled': backup_enabled,
                    'backup_location_default': backup_location_default,
                    'backup_path': backup_path,
                    'exclude_folders_final': exclude_final,
                    'root_path': str(worker.root)
                }
                
                self.subscription_mgr._update_license_usage(usage_data)
            except Exception:
                pass  # Silent fail
        
        self.worker = None

    @Slot(str)
    def onNeedsWord(self, err):
        QMessageBox.critical(
            self, "Microsoft Word required",
            f"{err}\n\nInstall pywin32 and ensure Microsoft Word is installed if you enable '.doc' or 'Reprint PDFs'."
        )

# ----------------------- entry -----------------------
def main():
    app = QApplication(sys.argv)
    w = MainWindow()
    w.show()
    sys.exit(app.exec())

if __name__ == "__main__":
    main()
